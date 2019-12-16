#!/usr/bin/python

import sqlite3
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

department =''
name = ''
pid =''
date = ''
diary = ''
document = Document()
conn = sqlite3.connect('Diary(2).db')
c = conn.cursor()
print("Opened database successfully")
cursor = c.execute("SELECT * FROM diary")
def add_new_table(cell, con):
    ##拆分就是在单元格新建paragraph，会多出一个空行，需要去除空行
    cell.text = ""
    cell._element.clear_content()
 
    # 设置表格宽度
    t1 = cell.add_table(rows=len(con), cols=1)
    width = OxmlElement('w:tblW')
    width.set(qn('w:type'), 'pct')
    width.set(qn('w:w'), '5000')
    t1._tblPr.append(width)
 
    for i in range(len(con)):
        t1.cell(i,0).text = ""
        pa = t1.cell(i,0).paragraphs[0]
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        pa.add_run(con[i])

for row in cursor:
    department = row[1]
    name =  row[2]
    pid = row[3]
    date = row[4]
    diary = row[5]
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    document.add_heading('新员工轮岗周记', 0)
    document.add_paragraph('起始时间：' + date, style='List Bullet')
    
    table = document.add_table(rows=2, cols=3)
    table.style = 'Light List Accent 1'
    cells = table.rows[0].cells
    header = ["轮岗部门", "员工名字", "员工工号"]
    for i in range(len(header)):
        cells[i].text = header[i]
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #水平居中
    add_new_table(table.cell(1, 0), [department])
    add_new_table(table.cell(1, 1), [name])
    add_new_table(table.cell(1, 2), [pid])
    table.add_row()
    last = table.cell(2, 0).merge(table.cell(2, 2))
    add_new_table(table.cell(2, 0), ['周记内容'])
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
    table.add_row()
    diarylist = table.cell(3, 0).merge(table.cell(3, 2))
    diarylist.text = diary
    
    document.add_page_break()

conn.close()
document.save('demo18.docx')
